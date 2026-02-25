import streamlit as st
import pandas as pd
import PyPDF2
import re
from io import BytesIO
import plotly.express as px
from datetime import datetime
import base64
import docx
import folium
from streamlit_folium import st_folium
import google.generativeai as genai

st.set_page_config(page_title="Certus Command Center", page_icon="🚂", layout="wide")

# --- AI CONFIGURATIE VOOR DE CHATBOT ---
API_KEY = "AIzaSyCiz1mWb378emBqRE3Tq3rLIIyFtm1fajI"
genai.configure(api_key=API_KEY)
# We gebruiken gemini-pro, dit is het meest stabiele model voor de chat
model = genai.GenerativeModel('gemini-pro') 

# --- DATABASE LOCATIES ---
LOCATIES_DB = {
    "GENT-ZEEH": [51.134, 3.823],
    "GENT-ZEEHAVEN": [51.134, 3.823],
    "VERB.GTS": [51.145, 3.815],
    "ANTWERPEN-NOORD": [51.275, 4.433],
    "ROOSENDAAL": [51.540, 4.458],
    "ZEEBRUGGE": [51.328, 3.197],
    "KALLO": [51.254, 4.275],
    "BRUSSEL-ZUID": [50.836, 4.335],
    "EVERGEM-BUNDEL ZANDEKEN": [51.121, 3.738],
    "EVERGEM": [51.108, 3.708]
}

def speel_certus_animatie():
    if 'animatie_gespeeld' not in st.session_state:
        try:
            with open("logo.png", "rb") as f:
                b64_logo = base64.b64encode(f.read()).decode("utf-8")
            css_animatie = f"""
            <style>
            #splash-screen {{ position: fixed; top: 0; left: 0; width: 100vw; height: 100vh; background-color: #0e1117; z-index: 99999; display: flex; justify-content: center; align-items: center; animation: fadeOut 1.5s forwards; animation-delay: 2s; pointer-events: none; }}
            #splash-logo {{ width: 350px; animation: moveAndShrink 1.5s forwards; animation-delay: 1.5s; }}
            @keyframes fadeOut {{ 0% {{ opacity: 1; }} 100% {{ opacity: 0; visibility: hidden; }} }}
            @keyframes moveAndShrink {{ 0% {{ transform: scale(1) translate(0, 0); opacity: 1; }} 100% {{ transform: scale(0.3) translate(-100vw, -100vh); opacity: 0; }} }}
            </style>
            <div id="splash-screen"><img id="splash-logo" src="data:image/png;base64,{b64_logo}"></div>
            """
            st.markdown(css_animatie, unsafe_allow_html=True)
            st.session_state.animatie_gespeeld = True
        except: pass

speel_certus_animatie()

if 'df_ritten' not in st.session_state:
    st.session_state.df_ritten = pd.DataFrame()

# Chatgeschiedenis initialiseren
if "messages" not in st.session_state:
    st.session_state.messages = []

def genereer_word_rapport(df):
    doc = docx.Document()
    titel = doc.add_heading('Certus Rail Solutions - Operationeel Rapport', 0)
    titel.alignment = 1 
    doc.add_paragraph(f"Gegenereerd op: {datetime.now().strftime('%d-%m-%Y %H:%M')}")
    doc.add_heading('1. Samenvatting van de Prestaties', level=1)
    tot_km = df['Afstand (km)'].sum()
    tot_ton = df['Gewicht (ton)'].sum()
    tot_ritten = len(df)
    p = doc.add_paragraph()
    p.add_run(f"In deze periode zijn er in totaal {tot_ritten} ritten uitgevoerd, met een totale afstand van {tot_km:,.1f} kilometer op het Infrabel netwerk. Het totale getransporteerde gewicht bedraagt {tot_ton:,.1f} ton.")
    doc.add_heading('2. Rittenspecificatie per Project', level=1)
    tabel = doc.add_table(rows=1, cols=6)
    tabel.style = 'Table Grid'
    hdr_cells = tabel.rows[0].cells
    for i, text in enumerate(['Datum', 'Project', 'Trein Nr.', 'Type Rit', 'Afstand (km)', 'Gewicht (ton)']):
        hdr_cells[i].text = text
    for index, row in df.iterrows():
        row_cells = tabel.add_row().cells
        row_cells[0].text, row_cells[1].text, row_cells[2].text = str(row['Datum']), str(row['Project']), str(row['Trein'])
        row_cells[3].text, row_cells[4].text, row_cells[5].text = str(row['Type']), f"{row['Afstand (km)']:.1f}", f"{row['Gewicht (ton)']:.1f}"
    f = BytesIO()
    doc.save(f)
    f.seek(0)
    return f

def analyseer_bestanden(files, gekozen_project):
    treinen = {}
    
    # VANGNET 1: Lees eerst de Excel bestandsnamen uit om bekende treinen te forceren
    bekende_treinen = set()
    for file in files:
        if file.name.lower().endswith(('.xlsx', '.xls')):
            t_nr_match = re.search(r'(\d{5})', file.name)
            if t_nr_match:
                bekende_treinen.add(t_nr_match.group(1))

    # PDF VERWERKING
    for file in files:
        try:
            if file.name.lower().endswith('.pdf'):
                reader = PyPDF2.PdfReader(file)
                text = "".join([page.extract_text() for page in reader.pages if page.extract_text()])
                if "infrabel" in text.lower() or "certus" in text.lower() or "trein" in text.lower():
                    datum_match = re.search(r'(\d{2})/(\d{2})/(\d{4})', text)
                    datum_str = f"{datum_match.group(3)}-{datum_match.group(2)}-{datum_match.group(1)}" if datum_match else datetime.today().strftime('%Y-%m-%d')
                    
                    # Zuig alle spaties en enters uit de tekst voor de moeilijke Infrabel formaten
                    clean_text = text.replace(' ', '').replace('\n', '')
                    ruwe_nummers = re.findall(r'([1-9]\d{4})(?=\d{2}/\d{2}/\d{2,4})', clean_text)
                    
                    # Oude regex als backup voor makkelijke formaten
                    oude_nummers = re.findall(r'(?<!\d)([1-9]\d{4})(?!\d)', text) + re.findall(r'([1-9]\d{4})(?=\s*\d{2}/\d{2})', text)
                    for n in oude_nummers:
                        if n not in ruwe_nummers:
                            ruwe_nummers.append(n)

                    # FORCEER DE BEKENDE TREINEN UIT DE EXCEL
                    for bt in bekende_treinen:
                        if bt in clean_text and bt not in ruwe_nummers:
                            ruwe_nummers.append(bt)
                    
                    un_codes = {'1202', '1863', '1965', '3257', '1203', '1170', '3082'}
                    nummers = list(set([n for n in ruwe_nummers if n not in un_codes]))
                    
                    km_match = re.search(r'(?:TreinKm|KmTrain|INFRABEL-net)[^\d]*(\d+(?:[.,]\d+)?)', text, re.IGNORECASE)
                    afstand = float(km_match.group(1).replace(',', '.')) if km_match else 0.0
                    
                    text_upper = text.upper()
                    route_match = re.search(r'([A-Z0-9.-]+)\s*->\s*([A-Z0-9.-]+)', text_upper)
                    
                    if route_match:
                        vertrek_loc = route_match.group(1)
                        aankomst_loc = route_match.group(2)
                    else:
                        vertrek_loc = "Onbekend"
                        aankomst_loc = "Onbekend"

                    for t_nr in nummers:
                        is_rid = "Ja" if re.search(r'RID:\s*Oui\s*/\s*Ja', text, re.IGNORECASE) or any(code in text for code in un_codes) else "Nee"
                        if t_nr in treinen:
                            treinen[t_nr].update({
                                "Afstand (km)": afstand, 
                                "Datum": datum_str,
                                "Vertrek": vertrek_loc,
                                "Aankomst": aankomst_loc
                            })
                            if is_rid == "Ja": treinen[t_nr]["RID"] = "Ja"
                        else:
                            treinen[t_nr] = {
                                "Datum": datum_str, "Project": gekozen_project, "Trein": t_nr, 
                                "Type": "Losse Rit", "Afstand (km)": afstand, "Gewicht (ton)": 0.0, 
                                "RID": is_rid, "UN": "",
                                "Vertrek": vertrek_loc, "Aankomst": aankomst_loc
                            }

            # EXCEL VERWERKING
            elif file.name.lower().endswith(('.xlsx', '.xls')):
                xl = pd.read_excel(file)
                if 'Trein' in xl.columns and 'Project' in xl.columns:
                    st.session_state.df_ritten = pd.concat([st.session_state.df_ritten, xl]).drop_duplicates(subset=['Trein'], keep='last')
                    continue
                t_nr_match = re.search(r'(\d{5})', file.name)
                if t_nr_match:
                    t_nr = t_nr_match.group(1)
                    excel_text = xl.to_string(index=False).upper()
                    un_match = re.search(r'\b(1202|1863|1965|3257|1203|1170|3082)\b', excel_text)
                    un_code = un_match.group(1) if un_match else ""
                    for col in xl.columns:
                        xl[col] = pd.to_numeric(xl[col].astype(str).str.replace(',', '.'), errors='coerce')
                    num_data = xl.select_dtypes(include=['number'])
                    zuivere_data = num_data[~num_data.isin([1202, 1863, 1965, 3257, 1203, 1170, 3082])]
                    gewicht = zuivere_data[zuivere_data < 4000].max().max()
                    gewicht = float(gewicht) if pd.notnull(gewicht) else 0.0
                    
                    type_rit = "Ledige Rit" if t_nr.endswith('1') or t_nr.endswith('3') or gewicht < 450 else "Beladen Rit"
                    
                    # VANGNET 2: Backup locaties uit de Excel als de PDF ze niet kon vinden!
                    gevonden_locs = [loc for loc in LOCATIES_DB.keys() if loc in excel_text]
                    v_loc = gevonden_locs[0] if gevonden_locs else "Onbekend"
                    a_loc = gevonden_locs[-1] if len(gevonden_locs) > 1 else v_loc

                    if t_nr in treinen: 
                        treinen[t_nr].update({"Gewicht (ton)": gewicht, "Type": type_rit})
                        if un_code: treinen[t_nr].update({"UN": un_code, "RID": "Ja"})
                        # Overschrijf 'Onbekend' met de Excel-locatie indien nodig
                        if treinen[t_nr].get("Vertrek", "Onbekend") == "Onbekend" and v_loc != "Onbekend":
                            treinen[t_nr]["Vertrek"] = v_loc
                            treinen[t_nr]["Aankomst"] = a_loc
                    else: 
                        treinen[t_nr] = {
                            "Datum": datetime.today().strftime('%Y-%m-%d'), "Project": gekozen_project, "Trein": t_nr, 
                            "Type": type_rit, "Afstand (km)": 0.0, "Gewicht (ton)": gewicht, 
                            "RID": "Ja" if un_code else "Nee", "UN": un_code,
                            "Vertrek": v_loc, "Aankomst": a_loc
                        }
        except Exception as e: 
            st.error(f"Fout bij {file.name}: {e}")
    return pd.DataFrame(list(treinen.values()))

with st.sidebar:
    st.write("🚂 **Certus Rail Solutions**")
    if st.button("🗑️ Wis Data & Chat"):
        st.session_state.df_ritten = pd.DataFrame()
        st.session_state.messages = []
        st.rerun()
    keuze = st.radio("Menu:", ("🏠 Home (Dashboard)", "📄 Invoer Ritten", "🖨️ Rapportage", "💬 AI Assistent"))

if keuze == "🏠 Home (Dashboard)":
    st.title("📊 Actueel Overzicht - Certus")
    df = st.session_state.df_ritten
    if not df.empty:
        c1, c2, c3, c4 = st.columns(4)
        c1.metric("Treinen", len(df))
        c2.metric("Km Met Last", f"{df[df['Type'] == 'Beladen Rit']['Afstand (km)'].sum():,.1f} km")
        c3.metric("Km Losse/Ledig", f"{df[df['Type'].isin(['Losse Rit', 'Ledige Rit'])]['Afstand (km)'].sum():,.1f} km")
        c4.metric("Ton", f"{df['Gewicht (ton)'].sum():,.1f} t")
        col_grafiek, col_kaart = st.columns([1.5, 1])
        with col_grafiek:
            st.plotly_chart(px.bar(df, x='Project', y='Afstand (km)', color='Type', barmode='group', template="plotly_dark"), use_container_width=True)
        with col_kaart:
            m = folium.Map(location=[51.05, 3.71], zoom_start=8, tiles="CartoDB dark_matter")
            
            if 'Vertrek' in df.columns and 'Aankomst' in df.columns:
                for index, row in df.iterrows():
                    vertrek_naam = str(row.get("Vertrek", "Onbekend"))
                    aankomst_naam = str(row.get("Aankomst", "Onbekend"))
                    
                    if vertrek_naam in LOCATIES_DB:
                        v_coords = LOCATIES_DB[vertrek_naam]
                        folium.Marker(
                            v_coords, 
                            popup=f"Vertrek Trein: {row['Trein']}<br>Locatie: {vertrek_naam}", 
                            icon=folium.Icon(color='green', icon='play', prefix='fa')
                        ).add_to(m)
                    
                    if aankomst_naam in LOCATIES_DB:
                        a_coords = LOCATIES_DB[aankomst_naam]
                        folium.Marker(
                            a_coords, 
                            popup=f"Aankomst Trein: {row['Trein']}<br>Locatie: {aankomst_naam}", 
                            icon=folium.Icon(color='red', icon='stop', prefix='fa')
                        ).add_to(m)

                    if vertrek_naam in LOCATIES_DB and aankomst_naam in LOCATIES_DB and vertrek_naam != aankomst_naam:
                        lijn_kleur = 'orange' if row['Type'] == 'Beladen Rit' else 'lightblue'
                        folium.PolyLine(
                            locations=[LOCATIES_DB[vertrek_naam], LOCATIES_DB[aankomst_naam]],
                            color=lijn_kleur,
                            weight=3,
                            opacity=0.8,
                            dash_array='5, 5' if row['Type'] != 'Beladen Rit' else None
                        ).add_to(m)

            st_folium(m, height=400, use_container_width=True)
    else: st.info("Geen data.")

elif keuze == "📄 Invoer Ritten":
    st.title("📄 Data Invoeren")
    gekozen_project = st.text_input("Projectnaam:", value="P419")
    uploaded_files = st.file_uploader("Upload bestanden", accept_multiple_files=True)
    if uploaded_files and st.button("🚀 Verwerk bestanden"):
        nieuw_df = analyseer_bestanden(uploaded_files, gekozen_project)
        st.session_state.df_ritten = pd.concat([st.session_state.df_ritten, nieuw_df]).drop_duplicates(subset=['Trein'], keep='last')
        st.rerun()
    if not st.session_state.df_ritten.empty:
        st.dataframe(st.session_state.df_ritten, use_container_width=True)

elif keuze == "🖨️ Rapportage":
    st.title("🖨️ Rapportage")
    if not st.session_state.df_ritten.empty:
        st.download_button("📄 Download Rapport", data=genereer_word_rapport(st.session_state.df_ritten), file_name="Rapport.docx")

# --- HET NIEUWE AI CHAT TABBLAD ---
elif keuze == "💬 AI Assistent":
    st.title("🤖 Certus AI Assistent")
    st.markdown("Stel mij vragen over de huidige operatie of de ingeladen ritten. Ik analyseer de actuele data direct voor je!")
    
    # Toon de chatgeschiedenis
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    # Chat input balk onderaan
    if prompt := st.chat_input("Bijv: 'Hoeveel ton is er vervoerd voor P419?'"):
        # Voeg gebruikersvraag toe aan geschiedenis
        st.session_state.messages.append({"role": "user", "content": prompt})
        
        with st.chat_message("user"):
            st.markdown(prompt)

        with st.chat_message("assistant"):
            # Converteer de huidige tabel naar een leesbaar formaat voor de AI
            df_context = "Er zijn nog geen ritten ingeladen in het systeem."
            if not st.session_state.df_ritten.empty:
                df_context = st.session_state.df_ritten.to_csv(index=False)
                
            # Vertel de AI wie hij is en geef hem de data
            systeem_prompt = f"""
            Je bent 'Certus AI', de slimme operationele assistent van Certus Rail Solutions (een dochteronderneming van Strukton België). 
            Jouw collega's zijn Jan en Klaas (directie) en Lana. Je spreekt professioneel, behulpzaam en altijd in het Nederlands.
            
            Hier is de actuele ritdata van vandaag in CSV-formaat:
            {df_context}
            
            Beantwoord de volgende vraag van de gebruiker accuraat op basis van de bovenstaande data. Als de data leeg is, vertel de gebruiker dan vriendelijk dat ze eerst bestanden moeten uploaden in het 'Invoer Ritten' tabblad.
            
            Vraag: {prompt}
            """
            
            try:
                # Roep de stabiele Gemini-Pro AI aan
                response = model.generate_content(systeem_prompt)
                full_response = response.text
                st.markdown(full_response)
                # Sla het antwoord op in de geschiedenis
                st.session_state.messages.append({"role": "assistant", "content": full_response})
            except Exception as e:
                st.error(f"Oeps, ik kon even geen verbinding maken met mijn brein. Foutmelding: {e}")
